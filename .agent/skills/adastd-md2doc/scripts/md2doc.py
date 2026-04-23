import os
import sys
import markdown
import requests
import base64
import json
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_font_th_sarabun(run):
    """Applies TH Sarabun New font to a run."""
    run.font.name = 'TH Sarabun New'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
    r.rPr.rFonts.set(qn('w:ascii'), 'TH Sarabun New')
    r.rPr.rFonts.set(qn('w:hAnsi'), 'TH Sarabun New')
    r.rPr.rFonts.set(qn('w:cs'), 'TH Sarabun New')

def set_table_borders(table):
    """Ensures table borders are visible."""
    tbl = table._tbl
    tblPr = tbl.xpath('w:tblPr')[0]
    tblBorders = OxmlElement('w:tblBorders')
    for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        node = OxmlElement(f'w:{border}')
        node.set(qn('w:val'), 'single')
        node.set(qn('w:sz'), '4')
        node.set(qn('w:space'), '0')
        node.set(qn('w:color'), 'auto')
        tblBorders.append(node)
    tblPr.append(tblBorders)

def get_mermaid_image(code):
    """Downloads mermaid image from mermaid.ink."""
    try:
        # Correct encoding for mermaid.ink: base64 of JSON object
        data = {"code": code, "mermaid": {"theme": "default"}}
        json_str = json.dumps(data)
        base64_bytes = base64.b64encode(json_str.encode('utf-8'))
        base64_string = base64_bytes.decode("ascii")
        url = "https://mermaid.ink/img/" + base64_string
        response = requests.get(url, timeout=15)
        if response.status_code == 200:
            return response.content
    except Exception as e:
        print(f"Error generating mermaid image: {e}")
    return None

def convert_md_to_docx(md_path, docx_path, template_path):
    # Load template
    if os.path.exists(template_path):
        doc = Document(template_path)
    else:
        doc = Document()

    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Convert MD to HTML
    html_content = markdown.markdown(md_content, extensions=['extra', 'tables', 'toc', 'fenced_code'])
    soup = BeautifulSoup(html_content, 'html.parser')

    for element in soup.contents:
        if element.name is None: continue

        # --- HEADINGS ---
        if element.name and len(element.name) == 2 and element.name.startswith('h') and element.name[1].isdigit():
            level = int(element.name[1])
            style = f'Heading {level}'
            try:
                p = doc.add_paragraph(element.get_text(), style=style)
            except:
                p = doc.add_paragraph(element.get_text())
                p.bold = True
            for run in p.runs:
                set_font_th_sarabun(run)

        # --- PARAGRAPHS ---
        elif element.name == 'p':
            p = doc.add_paragraph()
            for child in element.children:
                if child.name == 'strong':
                    run = p.add_run(child.get_text())
                    run.bold = True
                elif child.name == 'em':
                    run = p.add_run(child.get_text())
                    run.italic = True
                elif child.name == 'code':
                    run = p.add_run(child.get_text())
                    run.font.color.rgb = RGBColor(199, 37, 78) # Bootstrap-like code color
                else:
                    run = p.add_run(str(child))
                set_font_th_sarabun(run)

        # --- LISTS ---
        elif element.name in ['ul', 'ol']:
            for li in element.find_all('li'):
                style = 'List Bullet' if element.name == 'ul' else 'List Number'
                try:
                    p = doc.add_paragraph(li.get_text(), style=style)
                except:
                    p = doc.add_paragraph(f"• {li.get_text()}")
                for run in p.runs:
                    set_font_th_sarabun(run)
        
        # --- TABLES ---
        elif element.name == 'table':
            rows = element.find_all('tr')
            if not rows: continue
            
            max_cols = 0
            for row in rows:
                cols = row.find_all(['td', 'th'])
                max_cols = max(max_cols, len(cols))
            
            table = doc.add_table(rows=len(rows), cols=max_cols)
            set_table_borders(table)
            
            for i, row in enumerate(rows):
                cells = row.find_all(['td', 'th'])
                for j, cell in enumerate(cells):
                    if j >= max_cols: break
                    p = table.rows[i].cells[j].paragraphs[0]
                    p.text = cell.get_text()
                    if cell.name == 'th': # Header
                        for run in p.runs:
                            run.bold = True
                    for run in p.runs:
                        set_font_th_sarabun(run)

        # --- CODE BLOCKS (INCLUDING MERMAID) ---
        elif element.name == 'pre':
            code_tag = element.find('code')
            if code_tag:
                classes = code_tag.get('class', [])
                is_mermaid = 'language-mermaid' in classes
                code_text = code_tag.get_text()

                if is_mermaid:
                    # 1. Add Mermaid Image FIRST
                    img_data = get_mermaid_image(code_text)
                    if img_data:
                        import io
                        img_stream = io.BytesIO(img_data)
                        try:
                            # Add picture and center it (Reduced size to ~50% of page width)
                            pic = doc.add_picture(img_stream, width=Inches(3.0))
                            last_p = doc.paragraphs[-1]
                            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except Exception as e:
                            print(f"Error inserting picture: {e}")
                        
                        doc.add_paragraph() # Spacer

                    # 2. Add "Mermaid Graph Code:" Label
                    p_label = doc.add_paragraph()
                    run_label = p_label.add_run("Mermaid Graph Code:")
                    run_label.bold = True
                    set_font_th_sarabun(run_label)
                    
                    # 3. Add Mermaid Code Text
                    p_text = doc.add_paragraph()
                    for line in code_text.split('\n'):
                        run = p_text.add_run(line + '\n')
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
                else:
                    # Standard code block
                    p = doc.add_paragraph(code_text)
                    for run in p.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
                        set_font_th_sarabun(run) # Attempt to keep Thai chars readable

    # Global Font Fix
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            set_font_th_sarabun(run)

    doc.save(docx_path)
    print(f"Converted: {md_path} -> {docx_path}")

def main():
    if len(sys.argv) < 3:
        print("Usage: python md2doc.py <source_path> <target_dir>")
        sys.exit(1)

    source = sys.argv[1]
    target_dir = sys.argv[2]
    template = os.path.join(os.path.dirname(__file__), "..", "..", "..", "template", "adasoft-template.docx")

    if not os.path.exists(target_dir):
        os.makedirs(target_dir)

    if os.path.isfile(source):
        if source.endswith('.md'):
            filename = os.path.basename(source).replace('.md', '.docx')
            convert_md_to_docx(source, os.path.join(target_dir, filename), template)
    elif os.path.isdir(source):
        for root, dirs, files in os.walk(source):
            for file in files:
                if file.endswith('.md'):
                    rel_path = os.path.relpath(root, source)
                    dest_path = os.path.join(target_dir, rel_path)
                    if not os.path.exists(dest_path):
                        os.makedirs(dest_path)
                    
                    filename = file.replace('.md', '.docx')
                    convert_md_to_docx(os.path.join(root, file), os.path.join(dest_path, filename), template)

if __name__ == "__main__":
    main()
